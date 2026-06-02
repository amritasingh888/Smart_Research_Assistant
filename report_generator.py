import os
from datetime import datetime

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import (
    getSampleStyleSheet
)

from docx import Document


class ReportGenerator:

    def __init__(self):

        os.makedirs(
            "reports",
            exist_ok=True
        )

    # -----------------------------
    # TXT REPORT
    # -----------------------------

    def export_txt(
            self,
            content,
            filename="research_report.txt"):

        path = os.path.join(
            "reports",
            filename
        )

        with open(
            path,
            "w",
            encoding="utf-8"
        ) as f:

            f.write(content)

        return path

    # -----------------------------
    # PDF REPORT
    # -----------------------------

    def export_pdf(
            self,
            title,
            content,
            filename="research_report.pdf"):

        path = os.path.join(
            "reports",
            filename
        )

        doc = SimpleDocTemplate(
            path
        )

        styles = (
            getSampleStyleSheet()
        )

        story = []

        story.append(
            Paragraph(
                title,
                styles["Title"]
            )
        )

        story.append(
            Spacer(
                1,
                12
            )
        )

        story.append(
            Paragraph(
                content.replace(
                    "\n",
                    "<br/>"
                ),
                styles["BodyText"]
            )
        )

        doc.build(
            story
        )

        return path

    # -----------------------------
    # DOCX REPORT
    # -----------------------------

    def export_docx(
            self,
            title,
            content,
            filename="research_report.docx"):

        path = os.path.join(
            "reports",
            filename
        )

        doc = Document()

        doc.add_heading(
            title,
            level=1
        )

        doc.add_paragraph(
            content
        )

        doc.save(
            path
        )

        return path

    # -----------------------------
    # FULL RESEARCH REPORT
    # -----------------------------

    def create_full_report(
            self,
            summary,
            analysis,
            citations):

        report = f"""
SMART ACADEMIC RESEARCH ASSISTANT

Generated:
{datetime.now()}

==================================================
SUMMARY
==================================================

{summary}

==================================================
RESEARCH ANALYSIS
==================================================

Objectives
-----------
{analysis.get('objectives', '')}

Methodology
-----------
{analysis.get('methodology', '')}

Dataset
-----------
{analysis.get('dataset', '')}

Results
-----------
{analysis.get('results', '')}

Limitations
-----------
{analysis.get('limitations', '')}

Future Work
-----------
{analysis.get('future_work', '')}

==================================================
CITATIONS
==================================================

APA
-----------
{citations.get('APA', '')}

MLA
-----------
{citations.get('MLA', '')}

IEEE
-----------
{citations.get('IEEE', '')}

Chicago
-----------
{citations.get('Chicago', '')}
"""

        return report

    # -----------------------------
    # SAVE FULL REPORT
    # -----------------------------

    def export_complete_report(
            self,
            summary,
            analysis,
            citations):

        report = (
            self.create_full_report(
                summary,
                analysis,
                citations
            )
        )

        txt_file = self.export_txt(
            report,
            "complete_report.txt"
        )

        pdf_file = self.export_pdf(
            "Research Analysis Report",
            report,
            "complete_report.pdf"
        )

        docx_file = self.export_docx(
            "Research Analysis Report",
            report,
            "complete_report.docx"
        )

        return {

            "txt":
                txt_file,

            "pdf":
                pdf_file,

            "docx":
                docx_file
        }